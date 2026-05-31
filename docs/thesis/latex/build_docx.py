# -*- coding: utf-8 -*-
"""
Build a clean Word (.docx) version of the Ethnocount Bachelor's thesis from the
LaTeX sources in this folder.  No pandoc / LaTeX required -- pure python-docx.

TikZ diagrams cannot be rendered without LaTeX, so each figure is emitted as a
captioned placeholder that points to the PDF/LaTeX source for the rendered art.
Everything else -- headings, numbered/bulleted lists, description lists,
tables (incl. longtables), code listings, cross-references and the numbered
IEEE-style bibliography -- is converted faithfully.
"""
import io
import os
import re
import sys

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Mm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.abspath(__file__))
OUT = r"C:\Diplom\work\Ethnocount_Bachelor_Thesis_FINAL.docx"

# ----------------------------------------------------------------------------
# Thesis metadata (resolved from main.tex)
# ----------------------------------------------------------------------------
META = {
    "title_en": 'Design and Implementation of a Cross-Platform Treasury and '
                'Money-Transfer Platform ("Ethnocount") for Small '
                'Currency-Exchange Networks Using Flutter and Supabase',
    "title_lv": 'Starpplatformu kases un naudas pārvedumu platformas '
                '"Ethnocount" projektēšana un izstrāde nelielām valūtas maiņas '
                'tīkliem, izmantojot Flutter un Supabase',
    "student_en": "Farruh Muzrabov",
    "student_lv": "Farruhs Muzrabovs",
    "supervisor_en": "Dr.sc.ing., Prof. Jelena Chayko",
    "supervisor_lv": "Dr.sc.ing., prof. Jeļena Čaiko",
    "programme_en": "42484 Information Systems (BSc)",
    "programme_lv": "42484 Informācijas sistēmas (bak.)",
    "dept_en": "Department of Natural Sciences and Computer Engineering",
    "dept_lv": "Dabaszinātņu un datortehnoloģiju katedra",
    "year": "2026",
    "city_en": "Riga",
    "city_lv": "Rīga",
}

SCOPE_EN = ("The thesis comprises 74 pages, three chapters, 11 tables, "
            "8 figures, 12 listings, 2 appendices and 20 references.")
SCOPE_LV = ("Darba apjoms: 74 lappuses, trīs nodaļas, 11 tabulas, 8 attēli, "
            "12 listingi, 2 pielikumi un 20 literatūras avoti.")

BODY_FONT = "Times New Roman"
MONO_FONT = "Courier New"

# ----------------------------------------------------------------------------
# Cross-reference registry (label -> display string), built in pass 1.
# ----------------------------------------------------------------------------
REF = {
    "chap:litreview": "1",
    "chap:methodology": "2",
    "chap:implementation": "3",
    "chap:conclusions": "Conclusions",
    "chap:appendixA": "A",
    "chap:appendixB": "B",
}
# manual section map (chapter.section)
SEC = {
    # chapter 1
    "sec:domain": "1.1", "sec:doubleentry": "1.2", "sec:soa": "1.3",
    "sec:crossplatform": "1.4", "sec:baas": "1.5", "sec:cleanarch": "1.6",
    "sec:dsr": "1.7", "sec:gap": "1.8", "sec:ch1summary": "1.9",
    # chapter 2
    "sec:requirements": "2.1", "ssec:func": "2.1.1", "ssec:nonfunc": "2.1.2",
    "sec:domainmodel": "2.2", "sec:cleanarch2": "2.3", "sec:schema": "2.4",
    "sec:rls": "2.5", "sec:approval": "2.6", "sec:currency": "2.7",
    "sec:tradeoffs": "2.8", "sec:ch2summary": "2.9",
    # chapter 3
    "sec:stack": "3.1", "sec:transfer": "3.2", "sec:partner": "3.3",
    "sec:approvalimpl": "3.4", "sec:offline": "3.5", "sec:design": "3.6",
    "sec:testing": "3.7", "sec:perf": "3.8", "sec:uat": "3.9",
    "sec:limits": "3.10", "sec:ch3summary": "3.11",
    # appendix B
    "sec:appB-components": "B.1", "sec:appB-sequence": "B.2",
}
REF.update(SEC)

# Tables / figures / listings are numbered sequentially across the document
# (this matches the declared scope: 11 tables, 8 figures, 12 listings).
TAB = {
    "tab:taskmap": "1", "tab:soa": "2", "tab:frameworks": "3",
    "tab:func": "4", "tab:nonfunc": "5", "tab:tiers": "6", "tab:tradeoffs": "7",
    "tab:stack": "8", "tab:integration": "9", "tab:perf": "10", "tab:uat": "11",
}
FIG = {
    "fig:erd": "1", "fig:arch": "2", "fig:bloc": "3", "fig:approval": "4",
    "fig:approvalscreen": "5", "fig:testpyramid": "6",
    "fig:components": "7", "fig:sequence": "8",
}
LST = {
    "lst:transfers": "1", "lst:rls": "2", "lst:tree": "3", "lst:submit": "4",
    "lst:rpc": "5", "lst:bump": "6", "lst:request": "7",
    "lst:branches-ddl": "8", "lst:cp-ddl": "9", "lst:client-ddl": "10",
    "lst:transfers-ddl": "11", "lst:audit-ddl": "12",
}

CITE = {}  # filled after we sort the bibliography

# ----------------------------------------------------------------------------
# Inline LaTeX -> runs
# ----------------------------------------------------------------------------
SUP = str.maketrans("0123456789n+-=()", "⁰¹²³⁴⁵⁶⁷⁸⁹ⁿ⁺⁻⁼⁽⁾")
SUB = str.maketrans("0123456789", "₀₁₂₃₄₅₆₇₈₉")

ACCENTS = {
    '"o': "ö", '"a': "ä", '"u': "ü", '"O': "Ö", '"A': "Ä", '"U': "Ü",
    "'e": "é", "'a": "á", "'o": "ó", "'i": "í", "`a": "à", "`e": "è",
    "^o": "ô", "^e": "ê", "cs": "ş", "ug": "ğ", "cc": "ç",
}


def _math(s):
    """Convert a $...$ inner string to unicode-ish text."""
    s = s.replace("\\leq", "≤").replace("\\geq", "≥").replace("\\le", "≤")
    s = s.replace("\\approx", "≈").replace("\\times", "×").replace("\\neq", "≠")
    s = s.replace("\\to", "→").replace("\\cdot", "·").replace("\\,", " ")
    # superscripts ^{...} and ^x
    def sup(m):
        return m.group(1).translate(SUP)
    s = re.sub(r"\^\{([^}]*)\}", sup, s)
    s = re.sub(r"\^(.)", lambda m: m.group(1).translate(SUP), s)
    def sub(m):
        return m.group(1).translate(SUB)
    s = re.sub(r"_\{([^}]*)\}", sub, s)
    s = re.sub(r"_(.)", lambda m: m.group(1).translate(SUB), s)
    s = s.replace("{", "").replace("}", "")
    return s


def inline_runs(text):
    """Return list of (text, fmt) where fmt is a set of {'b','i','tt'}."""
    runs = []
    fmt = []          # stack of frames: each is set of flags added by that frame
    buf = []
    flags = set()

    def flush():
        if buf:
            runs.append(("".join(buf), frozenset(flags)))
            buf.clear()

    i, n = 0, len(text)
    while i < n:
        c = text[i]
        if c == "\\":
            # command
            m = re.match(r"\\(emph|textit|textbf|texttt|textsc|underline)\{", text[i:])
            if m:
                cmd = m.group(1)
                flush()
                fl = {"emph": "i", "textit": "i", "textbf": "b",
                      "texttt": "tt", "textsc": "", "underline": ""}[cmd]
                frame = set()
                if fl:
                    if fl not in flags:
                        frame.add(fl)
                    flags.add(fl)
                fmt.append(frame)
                i += m.end()
                continue
            m = re.match(r"\\(url)\{([^}]*)\}", text[i:])
            if m:
                flush()
                runs.append((m.group(2), frozenset(flags | {"url"})))
                i += m.end()
                continue
            m = re.match(r"\\(cite)\{([^}]*)\}", text[i:])
            if m:
                flush()
                keys = [k.strip() for k in m.group(2).split(",")]
                nums = [CITE.get(k, "?") for k in keys]
                buf.append("[" + "], [".join(str(x) for x in nums) + "]")
                i += m.end()
                continue
            m = re.match(r"\\(ref|cref|Cref|autoref)\{([^}]*)\}", text[i:])
            if m:
                cmd, label = m.group(1), m.group(2)
                disp = REF.get(label) or TAB.get(label) or FIG.get(label) or LST.get(label) or "?"
                if cmd in ("cref", "Cref", "autoref"):
                    pre = ("Table " if label in TAB else "Figure " if label in FIG
                           else "Listing " if label in LST
                           else "Section " if label.startswith("sec") or label.startswith("ssec")
                           else "Chapter " if label.startswith("chap") else "")
                    disp = pre + disp
                buf.append(disp)
                i += m.end()
                continue
            # accents like \"o  \'e  \c{s}  \u{g}  \"{o}
            m = re.match(r'\\(["\'`^~=.])\{?(\w)\}?', text[i:])
            if m:
                key = m.group(1) + m.group(2)
                buf.append(ACCENTS.get(key, m.group(2)))
                i += m.end()
                continue
            m = re.match(r"\\([cuvHbtdr])\{(\w)\}", text[i:])
            if m:
                key = m.group(1) + m.group(2)
                buf.append(ACCENTS.get(key, m.group(2)))
                i += m.end()
                continue
            # simple escaped chars
            m = re.match(r"\\([&%#_${}])", text[i:])
            if m:
                buf.append(m.group(1))
                i += m.end()
                continue
            m = re.match(r"\\(ldots|dots)", text[i:])
            if m:
                buf.append("…")
                i += m.end()
                continue
            m = re.match(r"\\,", text[i:])
            if m:
                buf.append(" ")
                i += 2
                continue
            m = re.match(r"\\ ", text[i:])
            if m:
                buf.append(" ")
                i += 2
                continue
            # drop other simple commands like \noindent \centering \singlespacing
            m = re.match(r"\\[a-zA-Z]+\*?", text[i:])
            if m:
                i += m.end()
                continue
            # stray backslash
            i += 1
            continue
        if c == "{":
            fmt.append(set())          # plain grouping frame
            i += 1
            continue
        if c == "}":
            flush()
            if fmt:
                frame = fmt.pop()
                for fl in frame:
                    flags.discard(fl)
            i += 1
            continue
        if c == "$":
            j = text.find("$", i + 1)
            if j == -1:
                j = n
            buf.append(_math(text[i + 1:j]))
            i = j + 1
            continue
        if c == "~":
            buf.append(" ")
            i += 1
            continue
        if c == "-" and text[i:i + 3] == "---":
            buf.append("—")
            i += 3
            continue
        if c == "-" and text[i:i + 2] == "--":
            buf.append("–")
            i += 2
            continue
        buf.append(c)
        i += 1
    flush()
    return runs


def add_runs(p, text, base_size=None):
    for txt, flags in inline_runs(text):
        if txt == "":
            continue
        r = p.add_run(txt)
        if "b" in flags:
            r.bold = True
        if "i" in flags:
            r.italic = True
        if "tt" in flags or "url" in flags:
            r.font.name = MONO_FONT
            r._element.rPr.rFonts.set(qn("w:cs"), MONO_FONT)
        if "url" in flags:
            r.font.color.rgb = RGBColor(0x1A, 0x4D, 0x8F)
            r.underline = True
        if base_size:
            r.font.size = Pt(base_size)
    return p


# ----------------------------------------------------------------------------
# docx helpers
# ----------------------------------------------------------------------------

def body_para(doc, text="", align="just", space_after=6, first_indent=True):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(space_after)
    pf.alignment = {"just": WD_ALIGN_PARAGRAPH.JUSTIFY,
                    "center": WD_ALIGN_PARAGRAPH.CENTER,
                    "left": WD_ALIGN_PARAGRAPH.LEFT}[align]
    if first_indent and align == "just":
        pf.first_line_indent = Cm(1.0)
    if text:
        add_runs(p, text)
    return p


def caption_para(doc, text, bold_label):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(bold_label + " ")
    r.bold = True
    r.font.size = Pt(11)
    # remaining caption text
    for txt, flags in inline_runs(text):
        rr = p.add_run(txt)
        rr.font.size = Pt(11)
        rr.italic = True
        if "tt" in flags:
            rr.font.name = MONO_FONT
    return p


def code_block(doc, code, caption_label, caption_text):
    caption_para(doc, caption_text, caption_label)
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    # shading
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "F4F5F7")
    cell._tc.get_or_add_tcPr().append(shd)
    _set_table_borders(tbl, "C0C4CC")
    _set_cell_margins(cell, left=60, right=60, top=40, bottom=40)
    cell.paragraphs[0]._p.getparent().remove(cell.paragraphs[0]._p)
    for ln in code.split("\n"):
        p = cell.add_paragraph()
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        r = p.add_run(ln if ln else " ")
        r.font.name = MONO_FONT
        r._element.rPr.rFonts.set(qn("w:cs"), MONO_FONT)
        r.font.size = Pt(12)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _set_cell_margins(cell, left=108, right=108, top=0, bottom=0):
    """Set per-cell margins in twips (1/20 pt). Tighter margins give code more
    usable width so 14pt monospace lines wrap less."""
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for edge, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:w"), str(val))
        e.set(qn("w:type"), "dxa")
        mar.append(e)
    tcPr.append(mar)


def _set_table_borders(tbl, color="999999"):
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), color)
        borders.append(e)
    tblPr.append(borders)


def figure_placeholder(doc, caption_label, caption_text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.columns[0].width = Cm(14)
    cell = tbl.cell(0, 0)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "EEF1F6")
    cell._tc.get_or_add_tcPr().append(shd)
    _set_table_borders(tbl, "B6BECF")
    p = cell.paragraphs[0]
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("[ Diagram — rendered in the PDF / LaTeX version of the thesis ]")
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x5A, 0x63, 0x73)
    caption_para(doc, caption_text, caption_label)


def heading(doc, text, level, numbered_prefix=None, page_break=False):
    if page_break and level == 1:
        doc.add_page_break()
    p = doc.add_paragraph()
    p.style = doc.styles[f"Heading {level}"]
    p.paragraph_format.keep_with_next = True
    if numbered_prefix:
        rp = p.add_run(numbered_prefix + "  ")
        rp.bold = True
    add_runs(p, text)
    for r in p.runs:
        r.font.name = BODY_FONT
        r.font.color.rgb = RGBColor(0x10, 0x14, 0x1C)
    return p


# ----------------------------------------------------------------------------
# Table parsing
# ----------------------------------------------------------------------------
def split_cells(row):
    parts = re.split(r"(?<!\\)&", row)
    return [c.strip() for c in parts]


def emit_table(doc, caption_label, caption_text, header, rows):
    caption_para(doc, caption_text, caption_label)
    ncol = max(len(header), max((len(r) for r in rows), default=1))
    tbl = doc.add_table(rows=0, cols=ncol)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header
    hr = tbl.add_row().cells
    for j in range(ncol):
        cell = hr[j]
        cell.paragraphs[0].text = ""
        txt = header[j] if j < len(header) else ""
        rr = cell.paragraphs[0]
        rr.paragraph_format.space_after = Pt(2)
        rr.paragraph_format.space_before = Pt(2)
        add_runs(rr, txt)
        for run in rr.runs:
            run.bold = True
            run.font.size = Pt(10.5)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), "E7EAF0")
        cell._tc.get_or_add_tcPr().append(shd)
    # body
    for row in rows:
        cells = tbl.add_row().cells
        for j in range(ncol):
            cell = cells[j]
            cell.paragraphs[0].text = ""
            txt = row[j] if j < len(row) else ""
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(2)
            para.paragraph_format.space_before = Pt(2)
            add_runs(para, txt)
            for run in para.runs:
                run.font.size = Pt(10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def parse_tabular_rows(block):
    """block = the inside of a tabular/longtable. Return (header, rows)."""
    # strip caption/label/rule/longtable-control lines
    block = re.sub(r"\\caption\{[^}]*\}", "", block)
    block = re.sub(r"\\label\{[^}]*\}", "", block)
    for cmd in (r"\toprule", r"\midrule", r"\bottomrule",
                r"\endfirsthead", r"\endhead", r"\endfoot"):
        block = block.replace(cmd, "")
    # split rows on \\
    raw = re.split(r"\\\\", block)
    rows = []
    for r in raw:
        r = re.sub(r"\[[0-9.]+\s*[a-z]*\]", "", r).strip()  # drop \\[2mm]
        if not r:
            continue
        rows.append(split_cells(r))
    # dedupe consecutive identical header rows (longtable firsthead+head)
    out = []
    for r in rows:
        if out and r == out[-1]:
            continue
        out.append(r)
    if not out:
        return [], []
    return out[0], out[1:]


# ----------------------------------------------------------------------------
# Generic block-level parser for a chapter body
# ----------------------------------------------------------------------------
def strip_comment(line):
    return re.sub(r"(?<!\\)%.*$", "", line)


def parse_body(doc, text, chap_num=None, in_appendix=False):
    """chap_num: '1','2','3','A','B' or None. Used for section numbering display."""
    lines = text.split("\n")
    i, n = 0, len(lines)
    sec_no = [0, 0]  # section, subsection counters

    def secnum(level):
        if chap_num is None:
            return None
        if level == 2:
            sec_no[0] += 1
            sec_no[1] = 0
            return f"{chap_num}.{sec_no[0]}"
        else:
            sec_no[1] += 1
            return f"{chap_num}.{sec_no[0]}.{sec_no[1]}"

    while i < n:
        line = lines[i]
        s = line.strip()
        sc = strip_comment(line).strip()

        if sc == "" and s == "":
            i += 1
            continue

        # --- environments -------------------------------------------------
        if sc.startswith("\\begin{lstlisting}"):
            opt = ""
            m = re.search(r"\[(.*)\]", sc)
            if m:
                opt = m.group(1)
            # caption/label may span; gather option text until closing ]
            cap = ""
            label = ""
            mc = re.search(r"caption=\{(.*?)\}\s*,?\s*label", opt)
            if mc:
                cap = mc.group(1)
            else:
                mc = re.search(r"caption=\{(.*?)\}\]?$", opt)
                if mc:
                    cap = mc.group(1)
            ml = re.search(r"label=\{([^}]*)\}", opt)
            if ml:
                label = ml.group(1)
            code = []
            i += 1
            while i < n and not lines[i].strip().startswith("\\end{lstlisting}"):
                code.append(lines[i])
                i += 1
            i += 1
            num = LST.get(label, "?")
            code_block(doc, "\n".join(code), f"Listing {num}.",
                       cap if cap else "")
            continue

        if sc.startswith("\\begin{table}") or sc.startswith("\\begin{longtable}") \
                or sc.startswith("{\\singlespacing"):
            # gather until matching \end{table}/\end{longtable}
            block = []
            depth_long = "longtable" in sc
            i += 1
            # we may have started on the {\singlespacing line; keep reading
            while i < n:
                l = lines[i]
                if l.strip().startswith("\\end{table}") or l.strip().startswith("\\end{longtable}"):
                    i += 1
                    # consume a trailing lone "}" from {\singlespacing ...}
                    if i < n and lines[i].strip() == "}":
                        i += 1
                    break
                block.append(l)
                i += 1
            btext = "\n".join(block)
            cap = ""
            mc = re.search(r"\\caption\{(.*?)\}", btext, re.S)
            if mc:
                cap = re.sub(r"\s+", " ", mc.group(1)).strip()
            label = ""
            ml = re.search(r"\\label\{([^}]*)\}", btext)
            if ml:
                label = ml.group(1)
            # extract tabular/longtable inner
            mt = re.search(r"\\begin\{(tabular|longtable)\}(\{[^}]*\}|\{.*?\})?(.*?)\\end\{\1\}",
                           btext, re.S)
            inner = mt.group(3) if mt else btext
            header, rows = parse_tabular_rows(inner)
            num = TAB.get(label, "?")
            emit_table(doc, f"Table {num}.", cap, header, rows)
            continue

        if sc.startswith("\\begin{figure}"):
            block = []
            i += 1
            while i < n and not lines[i].strip().startswith("\\end{figure}"):
                block.append(lines[i])
                i += 1
            i += 1
            btext = "\n".join(block)
            cap = ""
            mc = re.search(r"\\caption\{(.*?)\}\s*\\label", btext, re.S)
            if not mc:
                mc = re.search(r"\\caption\{(.*)\}", btext, re.S)
            if mc:
                cap = re.sub(r"\s+", " ", mc.group(1)).strip()
                cap = cap.rstrip("}")
            label = ""
            ml = re.search(r"\\label\{([^}]*)\}", btext)
            if ml:
                label = ml.group(1)
            num = FIG.get(label, "?")
            figure_placeholder(doc, f"Figure {num}.", cap)
            continue

        if sc.startswith("\\begin{enumerate}") or sc.startswith("\\begin{itemize}"):
            ordered = sc.startswith("\\begin{enumerate}")
            endtok = "\\end{enumerate}" if ordered else "\\end{itemize}"
            i += 1
            items = []
            cur = None
            while i < n and not lines[i].strip().startswith(endtok):
                ls = strip_comment(lines[i]).strip()
                if ls.startswith("\\item"):
                    if cur is not None:
                        items.append(cur)
                    cur = re.sub(r"^\\item\s*", "", ls)
                elif cur is not None:
                    cur += " " + ls
                i += 1
            if cur is not None:
                items.append(cur)
            i += 1
            for it in items:
                p = doc.add_paragraph(style="List Number" if ordered else "List Bullet")
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                p.paragraph_format.space_after = Pt(3)
                add_runs(p, it.strip())
            continue

        if sc.startswith("\\begin{description}"):
            i += 1
            items = []
            cur = None
            while i < n and not lines[i].strip().startswith("\\end{description}"):
                ls = strip_comment(lines[i]).strip()
                if ls.startswith("\\item"):
                    if cur is not None:
                        items.append(cur)
                    cur = re.sub(r"^\\item\s*", "", ls)
                elif cur is not None:
                    cur += " " + ls
                i += 1
            if cur is not None:
                items.append(cur)
            i += 1
            for it in items:
                m = re.match(r"\[(.*?)\]\s*(.*)", it.strip(), re.S)
                label = m.group(1) if m else ""
                rest = m.group(2) if m else it.strip()
                p = doc.add_paragraph()
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.left_indent = Cm(0.6)
                if label:
                    for txt, flags in inline_runs(label):
                        rr = p.add_run(txt + " ")
                        rr.bold = True
                        if "tt" in flags:
                            rr.font.name = MONO_FONT
                add_runs(p, rest)
            continue

        # --- headings -----------------------------------------------------
        m = re.match(r"\\chapter\*?\{(.*)\}", sc)
        if m:
            title = m.group(1)
            if in_appendix:
                heading(doc, title, 1, f"Appendix {chap_num}.", page_break=True)
            elif "*" in sc:
                heading(doc, title, 1, None, page_break=True)
            else:
                heading(doc, title, 1, f"{chap_num}.", page_break=True)
            i += 1
            continue
        m = re.match(r"\\(section|subsection)\*?\{(.*)\}", sc)
        if m:
            lvl = 2 if m.group(1) == "section" else 3
            starred = "*" in sc.split("{")[0]
            prefix = None if starred else secnum(lvl)
            heading(doc, m.group(2), lvl, prefix)
            i += 1
            continue
        m = re.match(r"\\paragraph\{(.*)\}(.*)", sc)
        if m:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.keep_with_next = True
            for txt, flags in inline_runs(m.group(1)):
                rr = p.add_run(txt)
                rr.bold = True
            rest = m.group(2).strip()
            # body that follows on subsequent lines becomes a normal paragraph
            buf = [rest] if rest else []
            i += 1
            while i < n and strip_comment(lines[i]).strip() != "" and \
                    not _is_block_start(strip_comment(lines[i]).strip()):
                buf.append(strip_comment(lines[i]).strip())
                i += 1
            if buf:
                body_para(doc, " ".join(buf).strip())
            continue

        # skip standalone label / format-only lines
        if re.match(r"^\\label\{", sc) or sc in ("}", "{"):
            i += 1
            continue

        # --- plain paragraph ---------------------------------------------
        buf = []
        while i < n:
            ls = strip_comment(lines[i]).strip()
            if ls == "":
                break
            if _is_block_start(ls) or re.match(r"\\(chapter|section|subsection|paragraph)\*?\{", ls):
                break
            buf.append(ls)
            i += 1
        if buf:
            body_para(doc, " ".join(buf).strip())
        else:
            i += 1


def _is_block_start(ls):
    return (ls.startswith("\\begin{lstlisting}") or ls.startswith("\\begin{table}")
            or ls.startswith("\\begin{longtable}") or ls.startswith("\\begin{figure}")
            or ls.startswith("\\begin{enumerate}") or ls.startswith("\\begin{itemize}")
            or ls.startswith("\\begin{description}") or ls.startswith("{\\singlespacing"))


# ----------------------------------------------------------------------------
# Bibliography
# ----------------------------------------------------------------------------
def parse_bib(path):
    txt = open(path, encoding="utf-8").read()
    entries = []
    for m in re.finditer(r"@(\w+)\{([^,]+),(.*?)\n\}", txt, re.S):
        etype, key, body = m.group(1), m.group(2).strip(), m.group(3)
        fields = {}
        for fm in re.finditer(r"(\w+)\s*=\s*\{(.*?)\}\s*,?\s*\n", body + "\n", re.S):
            fields[fm.group(1).lower()] = re.sub(r"\s+", " ", fm.group(2)).strip()
        entries.append((key, etype, fields))
    return entries


def author_sortkey(authorfield):
    # take first author's surname (before first comma or first word in braces)
    a = authorfield
    a = re.sub(r"\{(.*?)\}", r"\1", a)
    first = re.split(r"\band\b", a)[0].strip()
    if "," in first:
        sur = first.split(",")[0]
    else:
        sur = first.split()[-1] if first.split() else first
    return sur.lower()


def tex2txt(s):
    return "".join(t for t, _ in inline_runs(s)).strip()


def fmt_authors(authorfield):
    parts = [p.strip() for p in re.split(r"\s+and\s+", authorfield)]
    out = []
    for p in parts:
        t = tex2txt(p)
        out.append("et al." if t.lower() == "others" else t)
    return ", ".join(out)


def build_bibliography(doc, entries):
    # sort by surname, then title (approx IEEE nty)
    entries = sorted(entries, key=lambda e: (author_sortkey(e[2].get("author", "")),
                                             e[2].get("title", "").lower()))
    for idx, (key, etype, f) in enumerate(entries, 1):
        CITE[key] = idx
    heading(doc, "Literature", 1, None, page_break=True)
    for idx, (key, etype, f) in enumerate(entries, 1):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.first_line_indent = Cm(-0.8)
        rb = p.add_run(f"[{idx}] ")
        rb.bold = True
        authors = fmt_authors(f.get("author", ""))
        title = tex2txt(f.get("title", ""))
        year = f.get("year", "")
        if authors:
            r = p.add_run(authors + ", ")
        # title italic
        rt = p.add_run('"' + title + ',"')
        rt.italic = True
        tail = []
        for fld in ("booktitle", "journal", "publisher", "institution",
                    "organization", "howpublished", "address"):
            if f.get(fld):
                tail.append(tex2txt(f[fld]))
        extra = []
        if f.get("volume"):
            extra.append("vol. " + f["volume"])
        if f.get("number"):
            extra.append("no. " + f["number"])
        if f.get("pages"):
            extra.append("pp. " + f["pages"])
        if f.get("edition"):
            extra.append(f["edition"] + " ed.")
        tailtext = ""
        if tail:
            tailtext += " " + "; ".join(tail)
        if extra:
            tailtext += ", " + ", ".join(extra)
        if year:
            tailtext += ", " + year
        tailtext += "."
        add_runs(p, tailtext)


# ----------------------------------------------------------------------------
# Front matter
# ----------------------------------------------------------------------------
def hr_center(doc, text, size, bold=True, after=4, before=0):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    return p


LOGO_PATH = os.path.join(BASE, "a_frontmatter", "assets", "RNU_large_logo.png")


def title_page(doc, lang):
    en = lang == "en"

    # --- top-left wide RNU logo (52 mm, as in the .cls) ---
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(6)
    if os.path.exists(LOGO_PATH):
        p.add_run().add_picture(LOGO_PATH, width=Mm(52))

    # --- one line: left label, right programme (bold 12pt) ---
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.tab_stops.add_tab_stop(Mm(155), WD_TAB_ALIGNMENT.RIGHT)
    rl = p.add_run(("Study Programme" if en else "Studiju programma") + "\t" +
                   (META["programme_en"] if en else META["programme_lv"]))
    rl.bold = True
    rl.font.size = Pt(12)

    # --- centered department (bold 16pt) ---
    hr_center(doc, META["dept_en"] if en else META["dept_lv"], 16, before=10, after=12)

    # --- centered university name (Large) ---
    hr_center(doc, "RIGA NORDIC UNIVERSITY" if en else "RĪGAS ZIEMEĻVALSTU UNIVERSITĀTE",
              17, bold=False, before=8, after=1)
    hr_center(doc, "(RNU)", 17, bold=False, after=14)

    # --- thesis title (bold 16pt) + work type (bold 14pt) ---
    hr_center(doc, META["title_en"] if en else META["title_lv"], 16, after=12, before=4)
    hr_center(doc, "BACHELOR'S THESIS" if en else "BAKALAURA DARBS", 14, after=22)

    # --- student / supervisor block (left-aligned label + value) ---
    rows = ((("Student:" if en else "Students:"),
             META["student_en"] if en else META["student_lv"]),
            (("Supervisor:" if en else "Darba vadītājs:"),
             META["supervisor_en"] if en else META["supervisor_lv"]))
    tbl = doc.add_table(rows=0, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.columns[0].width = Cm(6)
    tbl.columns[1].width = Cm(8)
    for lbl, val in rows:
        cells = tbl.add_row().cells
        cells[0].width = Cm(6)
        cells[1].width = Cm(8)
        cells[0].paragraphs[0].add_run(lbl)
        cells[1].paragraphs[0].add_run(val)
        for c in cells:
            for r in c.paragraphs[0].runs:
                r.font.size = Pt(12)

    # --- bottom: city + year (push toward the bottom of the page) ---
    for _ in range(6):
        doc.add_paragraph()
    hr_center(doc, (META["city_en"] if en else META["city_lv"]) + " " + META["year"],
              12, bold=False, before=8)
    doc.add_page_break()


def add_toc(doc):
    heading(doc, "Table of Contents", 1, None, page_break=True)
    p = doc.add_paragraph()
    run = p.add_run()
    fldBegin = OxmlElement("w:fldChar")
    fldBegin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldSep = OxmlElement("w:fldChar")
    fldSep.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t")
    t.text = "Right-click and choose “Update Field” to build the table of contents."
    fldEnd = OxmlElement("w:fldChar")
    fldEnd.set(qn("w:fldCharType"), "end")
    run._r.append(fldBegin)
    run._r.append(instr)
    run._r.append(fldSep)
    run._r.append(t)
    run._r.append(fldEnd)


def unchapter_block(doc, title, paragraphs, page_break=True):
    heading(doc, title, 1, None, page_break=page_break)
    for para in paragraphs:
        body_para(doc, para)


# ----------------------------------------------------------------------------
# Styles & footer
# ----------------------------------------------------------------------------
def setup_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    for lvl, size in ((1, 16), (2, 14), (3, 12)):
        st = doc.styles[f"Heading {lvl}"]
        st.font.name = BODY_FONT
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0x10, 0x14, 0x1C)
        st.paragraph_format.space_before = Pt(12 if lvl == 1 else 8)
        st.paragraph_format.space_after = Pt(6)


def enable_update_fields(doc):
    """Tell Word to refresh all fields (the TOC, page numbers) when the
    document is first opened, so the table of contents builds itself."""
    settings = doc.settings.element
    upd = settings.find(qn("w:updateFields"))
    if upd is None:
        upd = OxmlElement("w:updateFields")
        settings.append(upd)
    upd.set(qn("w:val"), "true")


def add_page_numbers(doc):
    sec = doc.sections[0]
    footer = sec.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fb = OxmlElement("w:fldChar"); fb.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = "PAGE"
    fe = OxmlElement("w:fldChar"); fe.set(qn("w:fldCharType"), "end")
    run._r.append(fb); run._r.append(it); run._r.append(fe)


# ----------------------------------------------------------------------------
# MAIN
# ----------------------------------------------------------------------------
def read(p):
    return open(os.path.join(BASE, p), encoding="utf-8").read()


def extract_abstract(texpath, scope_text):
    raw = read(texpath)
    paras = []
    for block in raw.split("\n\n"):
        b = block.strip()
        if not b or b.startswith("\\frontmatterpage") or b.startswith("\\unchapter"):
            # may contain unchapter on its own line at top; strip those lines
            b = "\n".join(l for l in b.split("\n")
                          if not l.strip().startswith("\\frontmatterpage")
                          and not l.strip().startswith("\\unchapter"))
            b = b.strip()
            if not b:
                continue
        if b.startswith("\\thesisscope"):
            continue
        if b.startswith("\\textbf{"):
            continue  # keywords line handled separately
        paras.append(re.sub(r"\s+", " ", b))
    return paras


def main():
    doc = Document()
    # A4 + margins
    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21.0)
    sec.top_margin = Mm(25)
    sec.bottom_margin = Mm(25)
    sec.left_margin = Mm(35)
    sec.right_margin = Mm(20)
    setup_styles(doc)
    add_page_numbers(doc)
    enable_update_fields(doc)

    # bibliography first so CITE map is ready before any \cite is rendered
    bib_entries = parse_bib(os.path.join(BASE, "x_bibliography", "references.bib"))
    sorted_entries = sorted(bib_entries,
                            key=lambda e: (author_sortkey(e[2].get("author", "")),
                                           e[2].get("title", "").lower()))
    for idx, (key, etype, f) in enumerate(sorted_entries, 1):
        CITE[key] = idx

    # ---- title pages ----
    title_page(doc, "lv")
    title_page(doc, "en")

    # ---- abstracts ----
    lv = extract_abstract(os.path.join("a_frontmatter", "abstract_lv.tex"), SCOPE_LV)
    unchapter_block(doc, "Anotācija", lv + [SCOPE_LV])
    kw_lv = ("Atslēgvārdi: starpplatformu mobilā izstrāde, Flutter, BLoC, "
             "Supabase, PostgreSQL, Row-Level Security, glabātās procedūras, "
             "divkāršā ieraksta grāmatvedība, naudas pārvedumi, ārvalstu valūtas, "
             "vairāku filiāļu operācijas, vairāku nomnieku arhitektūra, idempotence.")
    kp = body_para(doc, "", first_indent=False)
    rkb = kp.add_run("Atslēgvārdi: "); rkb.bold = True
    kp.add_run(kw_lv.split(": ", 1)[1])

    en = extract_abstract(os.path.join("a_frontmatter", "abstract_en.tex"), SCOPE_EN)
    unchapter_block(doc, "Abstract", en + [SCOPE_EN])
    kp = body_para(doc, "", first_indent=False)
    rkb = kp.add_run("Keywords: "); rkb.bold = True
    kp.add_run("cross-platform mobile development, Flutter, BLoC, Supabase, "
               "PostgreSQL, Row-Level Security, stored procedures, double-entry "
               "accounting, money transfers, foreign exchange, multi-branch "
               "operations, multi-tenant architecture, idempotency.")

    # ---- keywords table (bilingual) ----
    kw_raw = read(os.path.join("a_frontmatter", "keywords.tex"))
    mt = re.search(r"\\begin\{tabular\}\{[^}]*\}(.*?)\\end\{tabular\}", kw_raw, re.S)
    if mt:
        header, rows = parse_tabular_rows(mt.group(1))
        heading(doc, "Key words / Atslēgvārdi", 1, None, page_break=True)
        # this table has no header row -> treat all as body, give it a header
        emit_table(doc, "Table.", "Key words (English / Latvian)",
                   ["English", "Latviski"], [header] + rows)

    # ---- TOC ----
    add_toc(doc)

    # ---- Introduction ----
    intro = read(os.path.join("b_chapters", "a_introduction", "introduction.tex"))
    intro = re.sub(r"^\s*\\unchapter\{[^}]*\}", "", intro, count=1)
    heading(doc, "Introduction", 1, None, page_break=True)
    parse_body(doc, intro, chap_num=None)

    # ---- Chapters 1-3 ----
    for num, path in (("1", ("b_chapters", "chapter1", "chapter1.tex")),
                      ("2", ("b_chapters", "chapter2", "chapter2.tex")),
                      ("3", ("b_chapters", "chapter3", "chapter3.tex"))):
        parse_body(doc, read(os.path.join(*path)), chap_num=num)

    # ---- Conclusions ----
    parse_body(doc, read(os.path.join("b_chapters", "x_conclusions", "conclusions.tex")),
               chap_num=None)

    # ---- Literature ----
    build_bibliography(doc, bib_entries)

    # ---- Appendices ----
    parse_body(doc, read(os.path.join("y_backmatter", "appendices", "appendixA.tex")),
               chap_num="A", in_appendix=True)
    parse_body(doc, read(os.path.join("y_backmatter", "appendices", "appendixB.tex")),
               chap_num="B", in_appendix=True)

    # ---- Acknowledgments ----
    ack = read(os.path.join("y_backmatter", "acknowledgments.tex"))
    ack_paras = []
    for b in ack.split("\n\n"):
        b = "\n".join(l for l in b.split("\n")
                      if not l.strip().startswith("\\frontmatterpage")
                      and not l.strip().startswith("\\unchapter")
                      and not l.strip().startswith("\\vspace"))
        b = re.sub(r"\s+", " ", b).strip()
        if b:
            ack_paras.append(b)
    unchapter_block(doc, "Acknowledgments / Pateicības", ack_paras)

    # ---- Declaration ----
    decl = read(os.path.join("y_backmatter", "declaration.tex"))
    heading(doc, "Affirmation / Apliecinājums", 1, None, page_break=True)
    body_para(doc, "Ar šo es, " + META["student_lv"] + ", apliecinu, ka bakalaura "
              "darbs ir izpildīts patstāvīgi, bez citu palīdzības, no svešiem "
              "avotiem ņemtie dati un definējumi ir uzrādīti darbā. Šis darbs "
              "nekādā veidā nav iesniegts nevienai citai pārbaudījuma komisijai "
              "un nekur nav publicēts.")
    body_para(doc, "Hereby I, " + META["student_en"] + ", affirm that the "
              "Bachelor's thesis was performed independently; sources of data "
              "and definitions are provided. This work has not been submitted to "
              "any other examination commission and has not been published "
              "elsewhere.")
    body_para(doc, "", first_indent=False)
    body_para(doc, META["city_lv"] + ", 20____ . ____________ .          "
              "______________________________", first_indent=False)
    sigp = body_para(doc, "", first_indent=False)
    sr = sigp.add_run("                                            "
                      "signature, name")
    sr.font.size = Pt(9)

    doc.save(OUT)
    print("Saved:", OUT)
    print("Cited/total references:", len(CITE))


if __name__ == "__main__":
    main()
